"""
DOCX parser.

Reads every paragraph and table cell from a .docx file and merges them
into a single clean text block.
"""

from app.ingestion import IngestionResult
from app.logging import get_logger

logger = get_logger(__name__)


def parse_docx(file_path: str, file_name: str, file_size_bytes: int, **kwargs) -> IngestionResult:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required to parse .docx files")

    try:
        doc = Document(file_path)
    except Exception as exc:
        logger.error("Failed to open DOCX file", exc_info=exc, path=file_path)
        return IngestionResult(
            text="",
            file_type="DOCX",
            file_name=file_name,
            file_size_bytes=file_size_bytes,
            pages=0,
            tickets_detected=0,
            warnings=[f"Failed to open document: {exc}"],
        )

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    merged = "\n".join(paragraphs)
    tickets = _count_tickets(merged)

    return IngestionResult(
        text=merged,
        file_type="DOCX",
        file_name=file_name,
        file_size_bytes=file_size_bytes,
        pages=len(doc.sections),
        tickets_detected=tickets,
    )


def _count_tickets(text: str) -> int:
    if not text.strip():
        return 0
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    return max(1, len(lines))
