from __future__ import annotations

import tempfile
from pathlib import Path
from unittest.mock import patch


class TestDocxParser:
    def _create_docx(self, paragraphs: list[str], tables: list[list[list[str]]] | None = None):
        from docx import Document

        doc = Document()
        for p in paragraphs:
            doc.add_paragraph(p)
        if tables:
            for table_data in tables:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                for i, row_data in enumerate(table_data):
                    for j, cell_data in enumerate(row_data):
                        table.cell(i, j).text = cell_data
        return doc

    def test_parse_docx_basic(self):
        doc = self._create_docx(["Hello World", "This is a test document"])
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            path = f.name

        try:
            from app.ingestion.docx_parser import parse_docx

            result = parse_docx(path, "test.docx", 100)
            assert "Hello World" in result.text
            assert result.tickets_detected >= 1
        finally:
            Path(path).unlink(missing_ok=True)

    def test_parse_docx_with_table(self):
        doc = self._create_docx(
            ["Summary"],
            tables=[[["Name", "Value"], ["Item1", "100"], ["Item2", "200"]]],
        )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            path = f.name

        try:
            from app.ingestion.docx_parser import parse_docx

            result = parse_docx(path, "test.docx", 100)
            assert "Summary" in result.text
            assert "Name | Value" in result.text
            assert "Item1 | 100" in result.text
        finally:
            Path(path).unlink(missing_ok=True)

    def test_parse_docx_empty(self):
        doc = self._create_docx([])
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            path = f.name

        try:
            from app.ingestion.docx_parser import parse_docx

            result = parse_docx(path, "test.docx", 100)
            assert result.text == ""
        finally:
            Path(path).unlink(missing_ok=True)

    def test_parse_docx_invalid_file(self):
        from app.ingestion.docx_parser import parse_docx

        result = parse_docx("/tmp/nonexistent_xyz.docx", "test.docx", 100)
        assert result.text == ""
        assert len(result.warnings) > 0

    def test_parse_docx_import_error(self):
        from app.ingestion.docx_parser import parse_docx

        with patch.dict("sys.modules", {"docx": None}):
            try:
                parse_docx("test.docx", "test.docx", 100)
            except ImportError as e:
                assert "python-docx" in str(e)

    def test_count_tickets_empty(self):
        from app.ingestion.docx_parser import _count_tickets

        assert _count_tickets("") == 0
        assert _count_tickets("   ") == 0

    def test_count_tickets_single(self):
        from app.ingestion.docx_parser import _count_tickets

        assert _count_tickets("Single line") == 1

    def test_count_tickets_multiple(self):
        from app.ingestion.docx_parser import _count_tickets

        assert _count_tickets("Line 1\nLine 2\nLine 3") == 3
